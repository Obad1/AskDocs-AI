"""DOCX document parser."""
from pathlib import Path
from typing import Dict
from docx import Document
from .document_parser import DocumentParser

class DOCXParser(DocumentParser):
    """Parser for DOCX files."""
    
    def parse(self, file_path: Path) -> Dict[str, str]:
        """Extract text from DOCX file."""
        text = ""
        metadata = {"paragraphs": 0}
        
        try:
            doc = Document(file_path)
            metadata["paragraphs"] = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n\n"
            
            if not text.strip():
                text = "❌ No readable text found in DOCX."
        except Exception as e:
            text = f"❌ Error extracting text from DOCX: {str(e)}"
        
        return {"text": text.strip(), "metadata": metadata}

